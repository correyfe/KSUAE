# import docx 
# doc = docx.Document()
# paragraph = doc.add_paragraph('это текст в стиле Quote')
# paragraph.style = 'Quote'
# paragraph2 = doc.add_paragraph('обычный текст')
# run = paragraph2.add_run(' и этот текст в стиле Title')
# run.style = 'Strong'

# doc.save('styled_document.docx')
# print('Документ сос стилями создан!')



# import docx

# doc = docx.Document('example.docx')
# for paragraph in doc.paragraphs:
#     paragraph.style = 'Norma1'

# doc.save(restyled.docx)



# import docx
# import os

# doc1 = docx.Document('example.docx')
# doc2 = docx.Document('restyled.docx')

# styles = []
# for paragraph in doc1.paragraphs:
#     styles.append(paragraph.style)

# for i in range(len(doc1.paragraphs)):
#     doc2.paragraphs[i].style = styles[i]

# doc2.save('restored.docx')    




# import docx
# doc = docx.Document('example.docx')
# doc.paragraphs[1].runs[0].style = 'Intense Emphasis'
# doc.paragraphs[1].runs[4].underline = True
# doc.save("restyled2.docx")





# import docx

# doc = docx.Document()
# doc.add_paragraph('Здравствуй, мир!')
# par1 = doc.add_paragraph('Это второй абзац')
# par2 = doc.add_paragraph('Это третий абзац')

# par1.add_run('Этот текст был добавлен во второй абзац')
# par2.add_run('Добавляем текст в третий абзац').bold = True

# doc.save('helloworld.docx')
# doc.add_paragraph('Здравствуй, мир!','Title')
# doc.save('helloworld.docx')



# import docx
# doc = docx.Document()
# doc.add_heading('Заголовок 0',0)
# doc.add_heading('Заголовок 1',1)
# doc.add_heading('Заголовок 2',2)
# doc.add_heading('Заголовок 3',3)
# doc.add_heading('Заголовок 4',4)

# doc.save('heading.docx')



# import docx

# doc = docx.Document()
# doc.add_paragraph('Это первая страница')
# doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
# doc.add_paragraph('Это вторая страница')
# doc.save('pages1.docx')


# import docx

# doc = docx.Document()
# table = doc.add_table(rows = 3, cols = 3)
# table.style = 'Table Grid'

# for row in range(3):
#     for col in range(3):
#         cell = table.cell(row,col)
#         cell.text = str(row+1)+str(col+1)
# doc.save('table1.docx')



import docx

doc = docx.Document()
doc.add_paragraph('Это первый абзац')
doc.add_picture('1.webp',width = docx.shared.Cm(10))
doc.save('picture1.docx')



import tkinter as tk
from tkinter import messagebox

def say_hello():
    name = entry.get()

    if name:
        greeting_labe1.config(text=f"Привет, {name}!")
    else:
        messagebox.showwarning("Внимание","Введите имя!")

root = tk.Tk()
root.title("Приветствие")
root.geometry("300x200")
labe1 = tk.labe1(root, text = "Введите ваше имя:")
labe1.pack(pady=10)
entry = tk.Entry(root)
labe1.pack(pady=5)
button = tk.Button(root, text="", font=("Arial", 12))
button.pack(pady=10)
greeting_labe1 = tk.labe1(root, text="", font=("Arial", 12))
greeting_labe1.pack(pady=10)
root.mainloop()


count = 0

count_labe1 = tk.labe1(root, text="Счет : 0", font=("Arial", 14))
count_labe1.pack(pady=20)

button_frame = tk.Frame(root)
button_frame.pack(pady=10)
inc_button = tk.Button(root, text="Нажми меня", command=increment, width=15)
inc_button.pack(side=tk.LEFT,)


